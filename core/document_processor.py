import os
import re
import PyPDF2
import uuid
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document
import streamlit as st
from config import PARENT_CHUNK_SIZE, PARENT_CHUNK_OVERLAP, CHILD_CHUNK_SIZE, CHILD_CHUNK_OVERLAP
import docx

def process_uploaded_files(uploaded_files):
    """
    Đọc các file được tải lên, làm sạch, và chia thành các parent và child chunks.
    Trả về một tuple: (parent_chunks, child_chunks)
    """
    print("[document_processor] Bắt đầu xử lý và chia chunk parent/child...")
    if not isinstance(uploaded_files, list):
        uploaded_files = [uploaded_files]
        print(f"[document_processor] Đã chuyển uploaded_files thành list, số lượng: 1")
    else:
        print(f"[document_processor] Số lượng file được tải lên: {len(uploaded_files)}")

    for i, uploaded_file in enumerate(uploaded_files):
        if uploaded_file is None:
            print(f"[document_processor] Warning: File #{i} là None")
        else:
            try:
                print(f"[document_processor] File #{i}: {uploaded_file.name}, size: {uploaded_file.size} bytes")
            except Exception as e:
                print(f"[document_processor] Error khi truy cập thông tin file #{i}: {e}")

    raw_docs_with_source = []
    for uploaded_file in uploaded_files:
        if uploaded_file is not None:
            text = ""
            try:
                file_name = uploaded_file.name
                file_extension = os.path.splitext(file_name)[1].lower()

                if file_extension == ".txt":
                    text = str(uploaded_file.read(), "utf-8")
                elif file_extension == ".pdf":
                    pdf_reader = PyPDF2.PdfReader(uploaded_file)
                    for page in pdf_reader.pages:
                        text += page.extract_text() or ""
                elif file_extension in [".doc", ".docx"]:
                    doc = docx.Document(uploaded_file)
                    text = "\n\n".join([paragraph.text for paragraph in doc.paragraphs])
                else:
                    st.error(f"Định dạng file '{file_name}' không được hỗ trợ.")
                    continue

                if not text.strip():
                    st.warning(f"File '{file_name}' không có nội dung.")
                    continue
                
                # Làm sạch văn bản
                text = re.sub(r'\s+', ' ', text).strip()
                raw_docs_with_source.append(Document(page_content=text, metadata={"source": file_name}))

            except Exception as e:
                st.error(f"Lỗi khi đọc file '{uploaded_file.name}': {e}")
    
    if not raw_docs_with_source:
        print("[document_processor] Không có văn bản nào được trích xuất.")
        return None, None
    
    parent_splitter = RecursiveCharacterTextSplitter(
        chunk_size=PARENT_CHUNK_SIZE, 
        chunk_overlap=PARENT_CHUNK_OVERLAP
    )
    parent_chunks = parent_splitter.split_documents(raw_docs_with_source)
    print(f"[document_processor] Đã chia thành {len(parent_chunks)} parent chunks.")

    id_key = "parent_id"
    for i, p_chunk in enumerate(parent_chunks):
        p_chunk.metadata[id_key] = str(uuid.uuid4())
        p_chunk.metadata["chunk_id"] = i

    child_splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHILD_CHUNK_SIZE, 
        chunk_overlap=CHILD_CHUNK_OVERLAP
    )
    
    child_chunks = []
    for p_chunk in parent_chunks:
        _child_docs_content = child_splitter.split_text(p_chunk.page_content)
        
        for j, _child_content in enumerate(_child_docs_content):
            child_metadata = p_chunk.metadata.copy()
            child_metadata["child_chunk_id"] = j
            child_doc = Document(page_content=_child_content, metadata=child_metadata)
            child_chunks.append(child_doc)

    print(f"[document_processor] Đã chia thành {len(child_chunks)} child chunks.")
    
    return parent_chunks, child_chunks

